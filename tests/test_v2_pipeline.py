from __future__ import annotations

import base64
from pathlib import Path
import sys
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from template_to_doc_v2.pipeline import TemplateToDocV2Pipeline
from template_to_doc_v2.llm import generate_local_text
from template_to_doc_v2.docx_replace import apply_replacements
from template_to_doc_v2.web import _extract_attachment_texts, _template_image_bytes


PNG_RED = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)
PNG_BLUE = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAIAAACQd1PeAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def _make_template(path: Path) -> None:
    document = Document()
    document.add_heading("1 Scope", level=1)
    document.add_paragraph("This template describes the original aircraft routing system.")
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Project"
    table.cell(0, 1).text = "{{project_name}}"
    table.cell(1, 0).text = "Summary"
    table.cell(1, 1).text = "{{summary}}"
    document.add_heading("2 Delivery", level=1)
    document.add_paragraph("Deliver one Word document and one verification report.")
    document.save(path)


def _make_heading_only_template(path: Path) -> None:
    document = Document()
    document.add_heading("1 Parent", level=1)
    document.add_heading("1.1 Child", level=2)
    document.add_paragraph("Child paragraph sample.")
    document.save(path)


def _make_blank_parent_template(path: Path) -> None:
    document = Document()
    document.add_heading("1 Parent", level=1)
    document.add_paragraph("")
    document.add_heading("1.1 Child", level=2)
    document.add_paragraph("Child paragraph sample.")
    document.save(path)


def _make_sow_like_template(path: Path) -> None:
    document = Document()
    document.add_paragraph("Background:")
    document.add_paragraph("Original background paragraph for the template.")
    document.add_paragraph("Work Packages: 3.5 Months")
    document.add_paragraph("WP1: 1.5 Month")
    numbered = document.add_paragraph("1. Low-altitude airspace definition and route planning;")
    numbered.paragraph_format.first_line_indent = Pt(24)
    document.add_paragraph("Goal: investigate the original AAM system.")
    document.add_paragraph("0571-28236568")
    document.add_paragraph("www.aossci.com")
    document.save(path)


def _make_image_template(path: Path) -> None:
    image = path.with_name("image.png")
    image.write_bytes(PNG_RED)
    document = Document()
    document.add_heading("1 Image Section", level=1)
    document.add_paragraph("Original image section text.")
    document.add_picture(str(image))
    document.add_paragraph("图1 原始图片")
    document.save(path)


def _make_header_image_template(path: Path) -> None:
    image = path.with_name("header_image.png")
    image.write_bytes(PNG_RED)
    document = Document()
    document.sections[0].header.paragraphs[0].add_run().add_picture(str(image))
    document.add_heading("1 Body", level=1)
    document.add_paragraph("Body text without inline images.")
    document.save(path)


def _make_cover_title_template(path: Path) -> None:
    document = Document()
    for _ in range(2):
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("智能航路航线规划关键技术研究")
        run.font.size = Pt(22)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("软件设计说明")
        run.font.size = Pt(20)

        date = document.add_paragraph()
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date.add_run("2026 年 03 月").font.size = Pt(14)
        document.add_page_break()
    document.add_heading("1 Scope", level=1)
    document.add_paragraph("Original body section.")
    document.save(path)


def _document_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as zf:
        return zf.read("word/document.xml").decode("utf-8")


def _zip_entry(path: Path, name: str) -> bytes:
    with zipfile.ZipFile(path) as zf:
        return zf.read(name)


def test_import_and_build_preserves_template_package_shape(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    _make_template(template)
    root = tmp_path / "v2"
    pipeline = TemplateToDocV2Pipeline(root)
    profile = pipeline.import_template(template, profile_id="demo")

    assert profile["profile_id"] == "demo"
    assert profile["stats"]["placeholder_count"] == 2
    assert profile["stats"]["section_target_count"] >= 1

    project = {
        "fields": {"project_name": "Collision Avoidance", "summary": "MARL-based system"},
        "sections": {
            profile["targets"]["sections"][0]["id"]: {
                "content": "This document describes a multi-agent reinforcement learning collision avoidance algorithm."
            }
        },
    }
    outputs = pipeline.build("demo", "case1", project=project)
    final_docx = Path(outputs["final_docx"])

    assert final_docx.exists()
    final_text = "\n".join(p.text for p in Document(final_docx).paragraphs)
    assert "multi-agent reinforcement learning" in final_text
    table_text = "\n".join(cell.text for table in Document(final_docx).tables for row in table.rows for cell in row.cells)
    assert "Collision Avoidance" in table_text
    assert "MARL-based system" in table_text
    assert len(Document(final_docx).tables) == len(Document(template).tables)

    xml = _document_xml(final_docx)
    assert "{{project_name}}" not in xml
    assert "{{summary}}" not in xml


def test_init_project_writes_editable_yaml(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    _make_template(template)
    root = tmp_path / "v2"
    pipeline = TemplateToDocV2Pipeline(root)
    pipeline.import_template(template, profile_id="demo")
    project = pipeline.init_project("demo", "case2")

    assert project["profile_id"] == "demo"
    assert (root / "projects" / "case2" / "project.yaml").exists()
    assert "sections" in project
    assert "global" in project
    assert {"topic", "background", "materials", "requirements"}.issubset(project["global"])


def test_import_detects_large_centered_cover_titles_as_fields(tmp_path: Path) -> None:
    template = tmp_path / "cover_title.docx"
    _make_cover_title_template(template)
    root = tmp_path / "v2"
    pipeline = TemplateToDocV2Pipeline(root)

    profile = pipeline.import_template(template, profile_id="cover", project_id="case")
    paragraph_fields = profile["targets"]["paragraph_fields"]
    by_label = {field["label"]: field for field in paragraph_fields}

    assert "智能航路航线规划关键技术研究" in by_label
    assert "软件设计说明" in by_label
    assert "2026 年 03 月" not in by_label
    assert len(by_label["智能航路航线规划关键技术研究"]["occurrences"]) == 2

    project = pipeline.init_project("cover", "case")
    project["fields"][by_label["智能航路航线规划关键技术研究"]["id"]] = "低空航路规划系统"
    project["fields"][by_label["软件设计说明"]["id"]] = "概要设计说明"
    outputs = pipeline.build("cover", "case", project=project)
    paragraphs = [p.text for p in Document(Path(outputs["final_docx"])).paragraphs]

    assert paragraphs.count("低空航路规划系统") == 2
    assert paragraphs.count("概要设计说明") == 2
    assert "智能航路航线规划关键技术研究" not in paragraphs
    assert outputs["replacement_report"]["paragraph_fields"] == 4


def test_build_can_generate_missing_section_with_provider_patch(
    tmp_path: Path,
    monkeypatch,
) -> None:
    template = tmp_path / "template.docx"
    _make_template(template)
    root = tmp_path / "v2"
    pipeline = TemplateToDocV2Pipeline(root)
    profile = pipeline.import_template(template, profile_id="demo")
    section_id = profile["targets"]["sections"][0]["id"]

    class FakeProvider:
        def generate(self, prompt: str, system: str, temperature: float = 0.2) -> str:
            return "Generated local section about cooperative collision avoidance."

    monkeypatch.setattr("template_to_doc_v2.pipeline.deepseek_provider", lambda *args, **kwargs: FakeProvider())

    project = {
        "fields": {"project_name": "Generated Case", "summary": "Generated Summary"},
        "sections": {section_id: {"prompt": "Use cooperative collision avoidance."}},
    }
    outputs = pipeline.build(
        "demo",
        "case3",
        project=project,
        generate=True,
        api_key="fake-key",
    )

    final_text = "\n".join(p.text for p in Document(Path(outputs["final_docx"])).paragraphs)
    assert "Generated local section about cooperative collision avoidance." in final_text


def test_heading_without_direct_body_gets_context_sample(tmp_path: Path) -> None:
    template = tmp_path / "heading_only.docx"
    _make_heading_only_template(template)
    root = tmp_path / "v2"
    pipeline = TemplateToDocV2Pipeline(root)
    profile = pipeline.import_template(template, profile_id="demo")

    child = profile["targets"]["sections"][0]
    assert child["title"] == "1.1 Child"
    assert "Child paragraph sample." in child["template_sample"]
    assert child["sample_source"] == "body"


def test_scan_skips_empty_structural_parent_heading(tmp_path: Path) -> None:
    template = tmp_path / "blank_parent.docx"
    _make_blank_parent_template(template)
    root = tmp_path / "v2"
    pipeline = TemplateToDocV2Pipeline(root)
    profile = pipeline.import_template(template, profile_id="demo")

    titles = [section["title"] for section in profile["targets"]["sections"]]
    assert "1 Parent" not in titles
    assert titles == ["1.1 Child"]


def test_old_wide_parent_section_does_not_override_child_section(tmp_path: Path) -> None:
    template = tmp_path / "blank_parent.docx"
    output = tmp_path / "output.docx"
    _make_blank_parent_template(template)
    profile = {
        "targets": {
            "table_fields": [],
            "placeholders": [],
            "sections": [
                {
                    "id": "parent",
                    "title": "1 Parent",
                    "level": 1,
                    "heading_paragraph_index": 0,
                    "body_start_paragraph_index": 1,
                    "body_end_paragraph_index": 4,
                    "path": "sections.parent.content",
                },
                {
                    "id": "child",
                    "title": "1.1 Child",
                    "level": 2,
                    "heading_paragraph_index": 2,
                    "body_start_paragraph_index": 3,
                    "body_end_paragraph_index": 4,
                    "path": "sections.child.content",
                },
            ],
        }
    }
    project = {
        "sections": {
            "parent": {"content": "Parent replacement should not be inserted."},
            "child": {"content": "Child replacement."},
        }
    }

    report = apply_replacements(template, output, profile, project)
    paragraphs = [p.text for p in Document(output).paragraphs if p.text.strip()]

    assert report.sections == 1
    assert "Parent replacement should not be inserted." not in paragraphs
    assert "1 Parent" in paragraphs
    assert "1.1 Child" in paragraphs
    assert "Child replacement." in paragraphs


def test_sow_like_scan_includes_front_matter_and_ignores_contact_items(tmp_path: Path) -> None:
    template = tmp_path / "sow_like.docx"
    _make_sow_like_template(template)
    root = tmp_path / "v2"
    pipeline = TemplateToDocV2Pipeline(root)
    profile = pipeline.import_template(template, profile_id="demo")

    titles = [section["title"] for section in profile["targets"]["sections"]]
    assert "Background:" in titles
    assert "WP1: 1.5 Month" in titles
    assert "0571-28236568" not in titles
    assert "www.aossci.com" not in titles
    assert "1. Low-altitude airspace definition and route planning;" not in titles


def test_image_targets_can_keep_or_replace_original_media(tmp_path: Path) -> None:
    template = tmp_path / "image_template.docx"
    _make_image_template(template)
    root = tmp_path / "v2"
    pipeline = TemplateToDocV2Pipeline(root)
    profile = pipeline.import_template(template, profile_id="demo")

    images = profile["targets"]["images"]
    assert profile["stats"]["image_count"] == 1
    assert images[0]["label"] == "图1 原始图片"
    assert images[0]["partname"].endswith(".png")

    project = pipeline.init_project("demo", "case_img")
    assert project["images"][images[0]["id"]]["mode"] == "keep"
    project["images"][images[0]["id"]] = {
        "mode": "replace",
        "file": {"name": "new.png", "data": base64.b64encode(PNG_BLUE).decode("ascii")},
    }
    output = pipeline.build("demo", "case_img", project=project)
    media_name = images[0]["partname"].lstrip("/")

    assert _zip_entry(Path(output["final_docx"]), media_name) == PNG_BLUE
    assert output["replacement_report"]["images"] == 1


def test_image_scan_detects_package_or_header_media(tmp_path: Path) -> None:
    template = tmp_path / "header_image_template.docx"
    _make_header_image_template(template)
    root = tmp_path / "v2"
    pipeline = TemplateToDocV2Pipeline(root)
    profile = pipeline.import_template(template, profile_id="demo")

    images = profile["targets"]["images"]
    assert profile["stats"]["image_count"] == 1
    assert images[0]["partname"].endswith(".png")
    assert images[0]["source_part"].startswith("/word/header")


def test_template_image_preview_reads_media_from_profile(tmp_path: Path) -> None:
    template = tmp_path / "image_template.docx"
    _make_image_template(template)
    root = tmp_path / "v2"
    pipeline = TemplateToDocV2Pipeline(root)
    profile = pipeline.import_template(template, profile_id="demo")
    image_id = profile["targets"]["images"][0]["id"]

    data, content_type = _template_image_bytes(root, pipeline, "demo", image_id)

    assert data == PNG_RED
    assert content_type == "image/png"


def test_section_replacement_strips_duplicate_heading_line(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    _make_heading_only_template(template)
    root = tmp_path / "v2"
    pipeline = TemplateToDocV2Pipeline(root)
    profile = pipeline.import_template(template, profile_id="demo")
    section = profile["targets"]["sections"][0]

    output = pipeline.build(
        "demo",
        "case_heading",
        project={"sections": {section["id"]: {"content": "1.1 Child\nNew body text."}}},
    )
    paragraphs = [p.text for p in Document(Path(output["final_docx"])).paragraphs if p.text.strip()]

    assert paragraphs.count("1.1 Child") == 1
    assert "New body text." in paragraphs


def test_web_extracts_text_attachment_for_section_generation() -> None:
    encoded = base64.b64encode("参考资料：核心是多智能体强化学习避撞算法。".encode("utf-8")).decode("ascii")

    extracted = _extract_attachment_texts([{"name": "notes.txt", "data": encoded}])

    assert extracted[0]["name"] == "notes.txt"
    assert "多智能体强化学习" in extracted[0]["text_excerpt"]


def test_llm_prompt_includes_section_attachments() -> None:
    class FakeProvider:
        prompt = ""

        def generate(self, prompt: str, system: str, temperature: float = 0.2) -> str:
            self.prompt = prompt
            return "generated"

    provider = FakeProvider()
    result = generate_local_text(
        prompt="仿写为避撞算法章节",
        template_sample="本软件主要针对城市低空复杂环境下的飞行器航路规划需求。",
        target_label="第一章 范围",
        project_context={"global": {"topic": "多智能体强化学习"}},
        attachments=[{"name": "notes.txt", "text_excerpt": "附件资料：MARL collision avoidance"}],
        provider=provider,
    )

    assert result == "generated"
    assert "附件资料" in provider.prompt
    assert "MARL collision avoidance" in provider.prompt


def test_llm_empty_local_prompt_uses_global_context() -> None:
    class FakeProvider:
        prompt = ""

        def generate(self, prompt: str, system: str, temperature: float = 0.2) -> str:
            self.prompt = prompt
            return "generated from global"

    provider = FakeProvider()
    result = generate_local_text(
        prompt="",
        template_sample="Template sample paragraph.",
        target_label="1 Scope",
        project_context={"global": {"topic": "global collision avoidance topic"}},
        provider=provider,
    )

    assert result == "generated from global"
    assert '"local_prompt": ""' in provider.prompt
    assert "global collision avoidance topic" in provider.prompt
    assert "local_prompt" in provider.prompt
    assert "template_reference" in provider.prompt


def test_llm_retries_when_output_is_only_heading() -> None:
    calls: list[str] = []

    class RetryProvider:
        def generate(self, prompt: str, system: str, temperature: float = 0.2) -> str:
            calls.append(prompt)
            if len(calls) == 1:
                return "1 Scope"
            return (
                "This section describes the new multi-agent reinforcement learning collision "
                "avoidance system, including its operating context, target users, primary "
                "functions, and expected document scope."
            )

    result = generate_local_text(
        prompt="",
        template_sample=(
            "This template section contains a long background paragraph about the original "
            "aircraft routing system, its operating conditions, expected users, system "
            "capabilities, constraints, and document scope. It is intentionally long enough "
            "to require paragraph-level expansion."
        ),
        target_label="1 Scope",
        project_context={"global": {"topic": "multi-agent reinforcement learning collision avoidance"}},
        provider=RetryProvider(),
    )

    assert len(calls) == 2
    assert "multi-agent reinforcement learning" in result
    assert "retry_reason" in calls[1]


def test_llm_prompt_includes_template_metrics_and_length_guard() -> None:
    class CaptureProvider:
        prompt = ""

        def generate(self, prompt: str, system: str, temperature: float = 0.2) -> str:
            self.prompt = prompt
            return "Generated paragraph that follows the template length and language."

    provider = CaptureProvider()
    generate_local_text(
        prompt="new topic",
        template_sample="目标：原模板目标。\n\n描述：原模板描述。",
        target_label="WP1",
        project_context={"global": {"topic": "topic"}},
        provider=provider,
    )

    assert "template_metrics" in provider.prompt
    assert "不要超过模板样例约 20%" in provider.prompt
    assert "不要主动新增模板" in provider.prompt
