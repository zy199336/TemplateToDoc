from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Any, Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.table import Table
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from .compat import ensure_legacy_imports


PLACEHOLDER_PATTERNS = [
    re.compile(r"\{\{\s*([A-Za-z0-9_.\-\u4e00-\u9fff]+)\s*\}\}"),
    re.compile(r"\[\s*\u5f85\u8865\u5145[:\uff1a]\s*([^\]\r\n]{1,40})\s*\]"),
    re.compile(r"\u3010\s*\u5f85\u8865\u5145[:\uff1a]\s*([^\u3011\r\n]{1,40})\s*\u3011"),
]


HEADING_RE = re.compile(
    r"^\s*((?:\d+(?:\.\d+)*|[一二三四五六七八九十]+|第[一二三四五六七八九十]+[章节篇])"
    r"[\s\.、．-]+.+|[A-Z][A-Za-z0-9 &/()_.:：-]{2,100})\s*$"
)
CONTACT_HEADING_RE = re.compile(
    r"^\s*(?:www\.|https?://|\S+@\S+|\+?\d[\d\s\-()]{5,})\s*$",
    re.I,
)
BLIP_TAG = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
EMBED_ATTR = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
LINK_ATTR = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link"
VML_IMAGEDATA_TAG = "{urn:schemas-microsoft-com:vml}imagedata"
REL_ID_ATTR = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
OFFICE_RELID_ATTR = "{urn:schemas-microsoft-com:office:office}relid"
IMAGE_RELTYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
MIN_MAJOR_TITLE_FONT_PT = 16.0
MAJOR_TITLE_SCAN_LIMIT = 80


@dataclass(frozen=True)
class ParagraphRef:
    paragraph: Paragraph
    scope: str
    index: int


@dataclass(frozen=True)
class BodyBlock:
    kind: str
    paragraph_index: int | None = None
    rows: int = 0
    cols: int = 0
    text: str = ""


def slugify(value: str, fallback: str) -> str:
    ascii_id = re.sub(r"[^0-9A-Za-z_]+", "_", str(value or "")).strip("_").lower()
    return ascii_id or fallback


def iter_body_paragraph_refs(document: DocumentObject) -> list[ParagraphRef]:
    return [ParagraphRef(paragraph, "body", index) for index, paragraph in enumerate(document.paragraphs)]


def body_blocks(document: DocumentObject) -> tuple[list[BodyBlock], dict[int, int]]:
    blocks: list[BodyBlock] = []
    paragraph_to_block: dict[int, int] = {}
    paragraph_index = 0
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph_to_block[paragraph_index] = len(blocks)
            blocks.append(BodyBlock("paragraph", paragraph_index=paragraph_index))
            paragraph_index += 1
        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            text = "\n".join(cell.text for row in table.rows for cell in row.cells)
            blocks.append(
                BodyBlock(
                    "table",
                    rows=len(table.rows),
                    cols=len(table.columns),
                    text=text,
                )
            )
    return blocks, paragraph_to_block


def iter_all_paragraphs(document: DocumentObject) -> Iterable[Paragraph]:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        seen_cells: set[int] = set()
        for row in table.rows:
            for cell in row.cells:
                yield from iter_cell_paragraphs(cell, seen_cells)
    for section in document.sections:
        for container in [
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ]:
            for paragraph in container.paragraphs:
                yield paragraph
            for table in container.tables:
                seen_cells = set()
                for row in table.rows:
                    for cell in row.cells:
                        yield from iter_cell_paragraphs(cell, seen_cells)


def iter_cell_paragraphs(cell: _Cell, seen_cells: set[int]) -> Iterable[Paragraph]:
    cell_key = id(cell._tc)
    if cell_key in seen_cells:
        return
    seen_cells.add(cell_key)
    for paragraph in cell.paragraphs:
        yield paragraph
    for table in cell.tables:
        for row in table.rows:
            for nested in row.cells:
                yield from iter_cell_paragraphs(nested, seen_cells)


def paragraph_outline_level(paragraph: Paragraph) -> int | None:
    p_pr = paragraph._p.pPr
    if p_pr is not None:
        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is not None:
            value = outline.get(qn("w:val"))
            if value is not None and value.isdigit():
                return int(value) + 1
    style = paragraph.style
    style_name = str(getattr(style, "name", "") or "")
    style_id = str(getattr(style, "style_id", "") or "")
    for value in (style_name, style_id):
        match = re.search(r"heading\s*(\d+)|标题\s*(\d+)", value, re.I)
        if match:
            return int(match.group(1) or match.group(2))
    return None


def looks_like_heading(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text or len(text) > 100:
        return False
    if CONTACT_HEADING_RE.match(text):
        return False
    if re.match(r"^\d+[\.．、]\s*\S+", text) and (
        paragraph.paragraph_format.first_line_indent is not None
        or paragraph.paragraph_format.left_indent is not None
    ):
        return False
    outline = paragraph_outline_level(paragraph)
    if outline is not None:
        return True
    if re.match(r"^[A-Z][A-Za-z0-9 &/()_.:：-]+[.。]$", text):
        return False
    return bool(HEADING_RE.match(text))


def heading_level(paragraph: Paragraph) -> int:
    outline = paragraph_outline_level(paragraph)
    if outline is not None:
        return outline
    text = paragraph.text.strip()
    if re.match(r"^WP\d+\b", text, re.I):
        return 2
    if re.match(r"^\d+\s+", text) or text.startswith("第"):
        return 1
    if re.match(r"^\d+\.\d+", text):
        return min(text.split()[0].count(".") + 1, 6)
    return 1


def contains_nontrivial_table(
    blocks: list[BodyBlock],
    paragraph_to_block: dict[int, int],
    heading_index: int,
    next_heading_index: int,
) -> bool:
    start_block = paragraph_to_block.get(heading_index, -1) + 1
    end_block = paragraph_to_block.get(next_heading_index, len(blocks))
    for block in blocks[start_block:end_block]:
        if block.kind != "table":
            continue
        if block.rows > 1 and block.cols > 1:
            return True
    return False


def paragraph_text_sample(paragraphs: list[Paragraph], start: int, end: int, limit: int = 1800) -> str:
    parts = [p.text.strip() for p in paragraphs[start:end] if p.text.strip()]
    text = "\n\n".join(parts)
    return text[:limit]


def section_context_sample(
    paragraphs: list[Paragraph],
    heading_index: int,
    start: int,
    end: int,
    limit: int = 1800,
) -> tuple[str, str]:
    direct = paragraph_text_sample(paragraphs, start, end, limit)
    if direct.strip():
        return direct, "body"
    parts: list[str] = []
    for paragraph in paragraphs[heading_index : min(len(paragraphs), heading_index + 12)]:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
        if len("\n\n".join(parts)) >= limit:
            break
    fallback = "\n\n".join(parts)[:limit]
    return fallback, "nearby_context" if fallback else "empty"


def scan_placeholders(document: DocumentObject) -> list[dict[str, Any]]:
    found: dict[str, dict[str, Any]] = {}
    occurrence_index = 0
    for paragraph in iter_all_paragraphs(document):
        text = paragraph.text
        for pattern in PLACEHOLDER_PATTERNS:
            for match in pattern.finditer(text):
                label = match.group(1).strip()
                token = match.group(0)
                field_id = slugify(label, f"placeholder_{occurrence_index + 1:02d}")
                occurrence_index += 1
                spec = found.setdefault(
                    field_id,
                    {
                        "id": field_id,
                        "label": label,
                        "kind": "placeholder",
                        "token": token,
                        "path": f"fields.{field_id}",
                        "default": "",
                        "occurrences": [],
                    },
                )
                spec["occurrences"].append({"token": token})
    return list(found.values())


def _clean_field_text(text: str) -> str:
    return " ".join(str(text or "").split())


def _paragraph_is_centered(paragraph: Paragraph) -> bool:
    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return True
    style = paragraph.style
    if style is not None:
        try:
            return style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            return False
    return False


def _run_font_size_pt(paragraph: Paragraph, run: Any) -> float | None:
    size = run.font.size
    if size is None and getattr(run, "style", None) is not None:
        size = run.style.font.size
    if size is None and paragraph.style is not None:
        size = paragraph.style.font.size
    return float(size.pt) if size is not None else None


def _paragraph_max_font_size_pt(paragraph: Paragraph) -> float | None:
    sizes = [
        size
        for run in paragraph.runs
        if run.text.strip()
        for size in [_run_font_size_pt(paragraph, run)]
        if size is not None
    ]
    if sizes:
        return max(sizes)
    if paragraph.style is not None and paragraph.style.font.size is not None:
        return float(paragraph.style.font.size.pt)
    return None


def _looks_like_major_title_field(paragraph: Paragraph, index: int) -> bool:
    text = _clean_field_text(paragraph.text)
    if not text or len(text) < 3 or len(text) > 80:
        return False
    if index >= MAJOR_TITLE_SCAN_LIMIT:
        return False
    if CONTACT_HEADING_RE.match(text):
        return False
    if paragraph_outline_level(paragraph) is not None:
        return False
    if str(getattr(paragraph.style, "name", "") or "").lower() == "caption":
        return False
    if re.search(r"(?:共\s*\d+\s*页|^\d{4}\s*年\s*\d{1,2}\s*月)", text):
        return False
    if not _paragraph_is_centered(paragraph):
        return False
    max_size = _paragraph_max_font_size_pt(paragraph)
    return max_size is not None and max_size >= MIN_MAJOR_TITLE_FONT_PT


def scan_paragraph_fields(document: DocumentObject) -> list[dict[str, Any]]:
    fields: list[dict[str, Any]] = []
    by_text: dict[str, dict[str, Any]] = {}
    for index, paragraph in enumerate(document.paragraphs):
        if not _looks_like_major_title_field(paragraph, index):
            continue
        text = _clean_field_text(paragraph.text)
        spec = by_text.get(text)
        if spec is None:
            field_id = slugify(text, f"paragraph_field_{len(fields) + 1:02d}")
            spec = {
                "id": field_id,
                "label": text,
                "kind": "paragraph_text",
                "default": text,
                "path": f"fields.{field_id}",
                "occurrences": [],
            }
            by_text[text] = spec
            fields.append(spec)
        spec["occurrences"].append({"paragraph_index": index})
    return fields


def _paragraph_image_rids(paragraph: Paragraph) -> list[str]:
    rids: list[str] = []
    for blip in paragraph._p.iter(BLIP_TAG):
        for attr in (EMBED_ATTR, LINK_ATTR):
            rid = blip.get(attr)
            if rid:
                rids.append(rid)
    for imagedata in paragraph._p.iter(VML_IMAGEDATA_TAG):
        for attr in (REL_ID_ATTR, OFFICE_RELID_ATTR):
            rid = imagedata.get(attr)
            if rid:
                rids.append(rid)
    return rids


def _next_caption(paragraphs: list[Paragraph], index: int) -> str:
    for paragraph in paragraphs[index + 1 : min(len(paragraphs), index + 5)]:
        text = paragraph.text.strip()
        if not text:
            continue
        if re.match(r"^(?:图|Figure|Fig\.?)\s*\d+", text, re.I):
            return text
        return ""
    return ""


def _image_part_from_relationship(rel: Any) -> Any | None:
    if rel is None or bool(getattr(rel, "is_external", False)):
        return None
    try:
        target_part = rel.target_part
    except Exception:
        return None
    partname = str(getattr(target_part, "partname", ""))
    content_type = str(getattr(target_part, "content_type", ""))
    reltype = str(getattr(rel, "reltype", ""))
    if reltype == IMAGE_RELTYPE or content_type.startswith("image/") or partname.startswith("/word/media/"):
        return target_part
    return None


def _iter_package_image_relationships(document: DocumentObject) -> Iterable[tuple[str, str, Any]]:
    for source_part in document.part.package.parts:
        source_partname = str(getattr(source_part, "partname", ""))
        for rid, rel in getattr(source_part, "rels", {}).items():
            target_part = _image_part_from_relationship(rel)
            if target_part is not None:
                yield source_partname, str(rid), target_part


def _unique_image_id(value: str, fallback: str, used_ids: set[str]) -> str:
    base = slugify(value, fallback)
    candidate = base
    suffix = 2
    while candidate in used_ids:
        candidate = f"{base}_{suffix}"
        suffix += 1
    used_ids.add(candidate)
    return candidate


def scan_images(document: DocumentObject) -> list[dict[str, Any]]:
    images: list[dict[str, Any]] = []
    occurrence = 0
    for index, paragraph in enumerate(document.paragraphs):
        for rid in _paragraph_image_rids(paragraph):
            part = paragraph.part.related_parts.get(rid)
            if part is None:
                continue
            occurrence += 1
            partname = str(getattr(part, "partname", ""))
            caption = _next_caption(document.paragraphs, index)
            image_id = slugify(caption or Path(partname).stem, f"image_{occurrence:02d}")
            images.append(
                {
                    "id": image_id,
                    "kind": "image",
                    "label": caption or f"图片 {occurrence}",
                    "caption": caption,
                    "paragraph_index": index,
                    "relationship_id": rid,
                    "partname": partname,
                    "content_type": str(getattr(part, "content_type", "")),
                    "path": f"images.{image_id}",
                    "default_mode": "keep",
                }
            )
    known_partnames = {str(image.get("partname") or "") for image in images}
    used_ids = {str(image.get("id") or "") for image in images}
    for source_partname, rid, part in _iter_package_image_relationships(document):
        partname = str(getattr(part, "partname", ""))
        if not partname or partname in known_partnames:
            continue
        occurrence += 1
        image_id = _unique_image_id(Path(partname).stem, f"image_{occurrence:02d}", used_ids)
        images.append(
            {
                "id": image_id,
                "kind": "image",
                "label": f"图片 {occurrence}",
                "caption": "",
                "paragraph_index": None,
                "relationship_id": rid,
                "source_part": source_partname,
                "partname": partname,
                "content_type": str(getattr(part, "content_type", "")),
                "path": f"images.{image_id}",
                "default_mode": "keep",
                "occurrences": [
                    {
                        "relationship_id": rid,
                        "source_part": source_partname,
                    }
                ],
            }
        )
        known_partnames.add(partname)
    return images


def scan_sections(document: DocumentObject) -> list[dict[str, Any]]:
    refs = iter_body_paragraph_refs(document)
    blocks, paragraph_to_block = body_blocks(document)
    headings = [(idx, ref.paragraph, heading_level(ref.paragraph)) for idx, ref in enumerate(refs) if looks_like_heading(ref.paragraph)]
    sections: list[dict[str, Any]] = []
    used_ids: set[str] = set()
    for order, (ref_index, paragraph, level) in enumerate(headings, start=1):
        next_ref_index = len(refs)
        first_child_ref_index: int | None = None
        for candidate_index, _candidate, candidate_level in headings[order:]:
            if candidate_level <= level:
                next_ref_index = candidate_index
                break
            if first_child_ref_index is None and candidate_level > level:
                first_child_ref_index = candidate_index
        title = paragraph.text.strip()
        base_id = slugify(title, f"section_{order:02d}")
        section_id = base_id
        suffix = 2
        while section_id in used_ids:
            section_id = f"{base_id}_{suffix}"
            suffix += 1
        used_ids.add(section_id)
        start = ref_index + 1
        end = first_child_ref_index or next_ref_index
        if start >= end:
            continue
        if contains_nontrivial_table(blocks, paragraph_to_block, ref_index, next_ref_index):
            continue
        body_paragraphs = [ref.paragraph for ref in refs]
        sample = paragraph_text_sample(body_paragraphs, start, end)
        if not sample.strip():
            continue
        sections.append(
            {
                "id": section_id,
                "title": title,
                "level": level,
                "kind": "section_text",
                "heading_paragraph_index": ref_index,
                "body_start_paragraph_index": start,
                "body_end_paragraph_index": end,
                "path": f"sections.{section_id}.content",
                "prompt_path": f"sections.{section_id}.prompt",
                "template_sample": sample,
                "sample_source": "body",
            }
        )
    return sections


def infer_table_fields(reference_docx: str | Path) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    ensure_legacy_imports()
    try:
        from template_to_doc.form_fields import infer_template_fields
    except Exception:
        return [], {}
    return infer_template_fields(reference_docx)


def inspect_template(reference_docx: str | Path) -> dict[str, Any]:
    document = Document(reference_docx)
    table_fields, project_fields = infer_table_fields(reference_docx)
    placeholders = scan_placeholders(document)
    paragraph_fields = scan_paragraph_fields(document)
    for placeholder in placeholders:
        project_fields.setdefault(
            placeholder["path"],
            {"label": placeholder["label"], "default": placeholder.get("default", "")},
        )
    for field in paragraph_fields:
        project_fields.setdefault(
            field["path"],
            {"label": field["label"], "default": field.get("default", "")},
        )
    sections = scan_sections(document)
    images = scan_images(document)
    return {
        "table_fields": table_fields,
        "placeholders": placeholders,
        "paragraph_fields": paragraph_fields,
        "sections": sections,
        "images": images,
        "project_fields": project_fields,
        "stats": {
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "sections": len(document.sections),
            "placeholder_count": len(placeholders),
            "paragraph_field_count": len(paragraph_fields),
            "table_field_count": len(table_fields),
            "section_target_count": len(sections),
            "image_count": len(images),
        },
    }
