from __future__ import annotations

import base64
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
import zipfile

from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph

from .compat import ensure_legacy_imports
from .docx_scan import heading_level, iter_all_paragraphs, looks_like_heading
from .yaml_io import dotted_get


@dataclass
class ReplacementReport:
    table_fields: int = 0
    paragraph_fields: int = 0
    placeholders: int = 0
    sections: int = 0
    images: int = 0
    unresolved: list[str] = field(default_factory=list)

    @property
    def total(self) -> int:
        return (
            self.table_fields
            + self.paragraph_fields
            + self.placeholders
            + self.sections
            + self.images
        )

    def as_dict(self) -> dict[str, Any]:
        return {
            "total": self.total,
            "table_fields": self.table_fields,
            "paragraph_fields": self.paragraph_fields,
            "placeholders": self.placeholders,
            "sections": self.sections,
            "images": self.images,
            "unresolved": sorted(set(self.unresolved)),
        }


def set_paragraph_text_preserving_runs(paragraph: Paragraph, value: str) -> bool:
    current = paragraph.text
    if current == value:
        return False
    if paragraph.runs:
        runs = paragraph.runs
        target_run = next((run for run in runs if run.text.strip()), runs[0])
        target_element = target_run._r
        target_run.text = value
        for run in runs:
            if run._r is not target_element:
                run.text = ""
    else:
        paragraph.add_run(value)
    return True


def replace_text_in_paragraph(paragraph: Paragraph, token: str, value: str) -> int:
    if token not in paragraph.text:
        return 0
    text = paragraph.text.replace(token, value)
    return 1 if set_paragraph_text_preserving_runs(paragraph, text) else 0


def fill_table_fields(document: DocumentObject, config: dict[str, Any], project: dict[str, Any]) -> int:
    if not config.get("template_fields"):
        return 0
    ensure_legacy_imports()
    try:
        from template_to_doc.docx.form_fill import TemplateFormFiller
    except Exception:
        return 0
    report = TemplateFormFiller(config).run(document, project)
    return int(report.replacements)


def fill_placeholders(
    document: DocumentObject,
    placeholders: list[dict[str, Any]],
    project: dict[str, Any],
) -> tuple[int, list[str]]:
    replacements = 0
    unresolved: list[str] = []
    for spec in placeholders:
        token = str(spec.get("token") or "")
        path = str(spec.get("path") or "")
        if not token or not path:
            continue
        value = dotted_get(project, path, spec.get("default", ""))
        if value in ("", None):
            unresolved.append(path)
            continue
        for paragraph in iter_all_paragraphs(document):
            replacements += replace_text_in_paragraph(paragraph, token, str(value))
    return replacements, unresolved


def fill_paragraph_fields(
    document: DocumentObject,
    paragraph_fields: list[dict[str, Any]],
    project: dict[str, Any],
) -> tuple[int, list[str]]:
    replacements = 0
    unresolved: list[str] = []
    for spec in paragraph_fields:
        path = str(spec.get("path") or "")
        if not path:
            continue
        value = dotted_get(project, path, spec.get("default", ""))
        if value in ("", None):
            unresolved.append(path)
            continue
        occurrences = spec.get("occurrences") or []
        for occurrence in occurrences:
            try:
                paragraph_index = int(occurrence.get("paragraph_index"))
            except (AttributeError, TypeError, ValueError):
                continue
            if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
                unresolved.append(f"{path}@{paragraph_index}")
                continue
            if set_paragraph_text_preserving_runs(document.paragraphs[paragraph_index], str(value)):
                replacements += 1
    return replacements, unresolved


def _remove_body_paragraph(document: DocumentObject, paragraph: Paragraph) -> None:
    element = paragraph._p
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _insert_paragraph_after(anchor: Paragraph, prototype: Paragraph, text: str) -> Paragraph:
    new_p = deepcopy(prototype._p)
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    set_paragraph_text_preserving_runs(paragraph, text)
    return paragraph


def _section_existing_text_paragraphs(
    document: DocumentObject,
    start: int,
    end: int,
    section_level: int | None = None,
) -> list[Paragraph]:
    candidates = document.paragraphs[start:end]
    paragraphs: list[Paragraph] = []
    for paragraph in candidates:
        if not paragraph.text.strip():
            continue
        if section_level is not None and looks_like_heading(paragraph):
            candidate_level = heading_level(paragraph)
            if candidate_level > section_level:
                break
        paragraphs.append(paragraph)
    return paragraphs


def _replacement_pieces(text: str) -> list[str]:
    normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        return []
    lines = [line.strip() for line in normalized.split("\n") if line.strip()]
    return lines or [normalized]


def _compact_heading(value: object) -> str:
    text = str(value or "")
    text = text.replace("\u3000", " ")
    text = "".join(text.split())
    return text.rstrip("：:；;。.、．")


def _strip_redundant_heading_pieces(pieces: list[str], section: dict[str, Any]) -> list[str]:
    target = _compact_heading(section.get("title") or section.get("id") or "")
    if not target:
        return pieces
    stripped = list(pieces)
    while stripped and _compact_heading(stripped[0]) == target:
        stripped.pop(0)
    return stripped


def replace_section_text(
    document: DocumentObject,
    section: dict[str, Any],
    value: str,
) -> bool:
    text = str(value or "").strip()
    if not text:
        return False
    start = int(section.get("body_start_paragraph_index", 0) or 0)
    end = int(section.get("body_end_paragraph_index", start) or start)
    pieces = _replacement_pieces(text)
    pieces = _strip_redundant_heading_pieces(pieces, section)
    if not pieces:
        return False
    section_level = int(section.get("level", 0) or 0) or None
    paragraphs = _section_existing_text_paragraphs(document, start, end, section_level)
    if not paragraphs:
        return False

    anchor = paragraphs[0]
    for index, piece in enumerate(pieces):
        if index < len(paragraphs):
            paragraph = paragraphs[index]
            set_paragraph_text_preserving_runs(paragraph, piece)
            anchor = paragraph
        else:
            prototype = paragraphs[-1]
            anchor = _insert_paragraph_after(anchor, prototype, piece)
    for old in paragraphs[len(pieces) :]:
        _remove_body_paragraph(document, old)
    return True


def _decode_uploaded_file(payload: dict[str, Any]) -> bytes:
    encoded = str(payload.get("data") or "")
    if not encoded:
        return b""
    return base64.b64decode(encoded)


def _normalized_image_ext(value: str) -> str:
    suffix = Path(value).suffix.lower()
    if suffix == ".jpeg":
        return ".jpg"
    return suffix


def _image_payload_is_compatible(partname: str, upload_name: str) -> bool:
    template_ext = _normalized_image_ext(partname)
    upload_ext = _normalized_image_ext(upload_name)
    if not template_ext or not upload_ext:
        return True
    return template_ext == upload_ext


def replace_images_in_package(
    output_docx: Path,
    images: list[dict[str, Any]],
    project: dict[str, Any],
) -> tuple[int, list[str]]:
    configured = project.get("images", {})
    if not isinstance(configured, dict) or not images:
        return 0, []
    replacements: dict[str, bytes] = {}
    unresolved: list[str] = []
    for image in images:
        image_id = str(image.get("id") or "")
        state = configured.get(image_id, {})
        if not isinstance(state, dict) or state.get("mode") != "replace":
            continue
        file_payload = state.get("file")
        if not isinstance(file_payload, dict):
            unresolved.append(f"images.{image_id}.file")
            continue
        partname = str(image.get("partname") or "").lstrip("/")
        upload_name = str(file_payload.get("name") or "")
        if not partname:
            unresolved.append(f"images.{image_id}.partname")
            continue
        if not _image_payload_is_compatible(partname, upload_name):
            unresolved.append(
                f"images.{image_id}: uploaded image extension must match template image extension ({Path(partname).suffix})"
            )
            continue
        data = _decode_uploaded_file(file_payload)
        if not data:
            unresolved.append(f"images.{image_id}.file")
            continue
        replacements[partname] = data
    if not replacements:
        return 0, unresolved

    temp = output_docx.with_suffix(output_docx.suffix + ".tmp")
    with zipfile.ZipFile(output_docx, "r") as source, zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = replacements.get(item.filename)
            if data is None:
                data = source.read(item.filename)
            target.writestr(item, data)
    temp.replace(output_docx)
    return len(replacements), unresolved


def fill_sections(
    document: DocumentObject,
    sections: list[dict[str, Any]],
    project: dict[str, Any],
) -> tuple[int, list[str]]:
    replacements = 0
    unresolved: list[str] = []
    # Replace from back to front so paragraph indexes remain valid.
    ordered = sorted(
        sections,
        key=lambda item: int(item.get("heading_paragraph_index", 0) or 0),
        reverse=True,
    )
    for section in ordered:
        path = str(section.get("path") or "")
        value = dotted_get(project, path, "")
        if value in ("", None):
            continue
        if replace_section_text(document, section, str(value)):
            replacements += 1
        else:
            unresolved.append(path)
    return replacements, unresolved


def apply_replacements(
    input_docx: str | Path,
    output_docx: str | Path,
    profile: dict[str, Any],
    project: dict[str, Any],
) -> ReplacementReport:
    document = Document(input_docx)
    report = ReplacementReport()
    config = {"template_fields": profile.get("targets", {}).get("table_fields", [])}
    report.table_fields = fill_table_fields(document, config, project)
    paragraph_fields = profile.get("targets", {}).get("paragraph_fields", [])
    if isinstance(paragraph_fields, list):
        report.paragraph_fields, unresolved = fill_paragraph_fields(document, paragraph_fields, project)
        report.unresolved.extend(unresolved)
    placeholders = profile.get("targets", {}).get("placeholders", [])
    if isinstance(placeholders, list):
        report.placeholders, unresolved = fill_placeholders(document, placeholders, project)
        report.unresolved.extend(unresolved)
    sections = profile.get("targets", {}).get("sections", [])
    if isinstance(sections, list):
        report.sections, unresolved = fill_sections(document, sections, project)
        report.unresolved.extend(unresolved)
    target = Path(output_docx)
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)
    images = profile.get("targets", {}).get("images", [])
    if isinstance(images, list):
        report.images, unresolved = replace_images_in_package(target, images, project)
        report.unresolved.extend(unresolved)
    return report
