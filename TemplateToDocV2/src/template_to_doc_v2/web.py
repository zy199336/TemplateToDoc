from __future__ import annotations

import base64
from html import escape
from io import BytesIO
import json
import shutil
import subprocess
import tempfile
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, quote, unquote, urlparse
from typing import Any
import zipfile

from .llm import deepseek_provider, generate_local_text
from .pipeline import TemplateToDocV2Pipeline


DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _json_bytes(data: dict[str, Any]) -> bytes:
    return json.dumps(data, ensure_ascii=False, default=str).encode("utf-8")


def _read_json(handler: BaseHTTPRequestHandler) -> dict[str, Any]:
    length = int(handler.headers.get("content-length", "0"))
    payload = json.loads(handler.rfile.read(length) or b"{}")
    if not isinstance(payload, dict):
        return {}
    return payload


def _read_default_deepseek_key(root: Path) -> str:
    key_path = root / "deepseekkey.txt"
    if not key_path.exists():
        return ""
    return key_path.read_text(encoding="utf-8", errors="replace").strip()


def _html(default_api_key: str = "") -> bytes:
    escaped_key = escape(default_api_key, quote=True)
    template = r"""<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>TemplateToDocV2</title>
<style>
body { font-family: system-ui, "Microsoft YaHei", sans-serif; margin: 0; color: #0b1f35; background: #f6f8f4; }
main { max-width: 1320px; margin: 0 auto; padding: 24px; }
section { border: 1px solid #cfd8c8; background: #fff; border-radius: 8px; padding: 18px; margin: 14px 0; }
h1 { font-size: 28px; margin: 0 0 12px; }
h2 { font-size: 20px; margin: 0 0 14px; }
label { display: block; font-weight: 700; margin: 10px 0 6px; }
input, textarea, select { width: 100%; box-sizing: border-box; border: 1px solid #b9c6b1; border-radius: 6px; padding: 9px; font: inherit; background: #fff; }
textarea { min-height: 120px; resize: vertical; }
button, .button { display: inline-flex; align-items: center; gap: 6px; border: 0; border-radius: 7px; padding: 10px 16px; color: #fff; background: #08756c; font-weight: 800; cursor: pointer; text-decoration: none; }
button.secondary { background: #2a405d; }
button.ghost { background: #61736a; }
.row { display: grid; grid-template-columns: repeat(3, 1fr); gap: 14px; }
.global-grid { display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 14px; }
.target { display: grid; grid-template-columns: minmax(0, 1fr) minmax(320px, 0.8fr); gap: 18px; align-items: start; }
.sample { border: 1px solid #d8dfd2; background: #fbfcfa; border-radius: 8px; padding: 12px; max-height: 360px; overflow: auto; white-space: pre-wrap; line-height: 1.55; }
.actions { display: flex; flex-wrap: wrap; gap: 10px; margin-top: 12px; }
.section-status { margin-top: 10px; min-height: 24px; color: #075f58; font-weight: 700; white-space: pre-wrap; }
.section-status.error { color: #9b1c1c; }
.global-status { margin-top: 10px; min-height: 24px; color: #075f58; font-weight: 700; white-space: pre-wrap; }
.global-status.error { color: #9b1c1c; }
.file-row { margin-top: 10px; }
.image-card { border: 1px solid #9bb8d8; border-radius: 8px; padding: 14px; margin: 12px 0; background: #f7fbff; }
.image-layout { display: grid; grid-template-columns: minmax(0, 1fr) minmax(360px, 0.9fr); gap: 18px; align-items: start; }
.image-preview-grid { display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 12px; }
.preview-box { border: 1px solid #b7cbe3; background: #fff; border-radius: 8px; padding: 10px; min-height: 180px; display: flex; flex-direction: column; gap: 8px; }
.preview-title { font-weight: 800; color: #17395f; }
.preview-frame { flex: 1; min-height: 130px; border: 1px dashed #b7cbe3; border-radius: 6px; display: flex; align-items: center; justify-content: center; overflow: hidden; background: #f9fbff; color: #5f6f7e; text-align: center; padding: 8px; }
.preview-frame img { max-width: 100%; max-height: 220px; object-fit: contain; }
.status { white-space: pre-wrap; color: #075f58; font-weight: 700; }
.muted { color: #5f6f7e; font-size: 14px; }
@media (max-width: 900px) { .row, .global-grid, .target, .image-layout, .image-preview-grid { grid-template-columns: 1fr; } }
</style>
</head>
<body>
<main>
<h1>TemplateToDocV2</h1>
<p class="muted">模板原件保留 + 局部文字替换/局部生成。最终 Word 从原模板复制后局部写回，不走整篇 Markdown 重建。</p>

<section>
<h2>1. 导入模板</h2>
<div class="row">
  <div><label>模板路径</label><input id="templatePath" placeholder="例如 TemplateToDoc/模板2.doc"></div>
  <div><label>或上传 Word</label><input id="templateFile" type="file" accept=".doc,.docx"></div>
  <div><label>Profile ID</label><input id="profileId" placeholder="自动使用文件名"></div>
</div>
<label>Project ID</label><input id="projectId" value="demo">
<div class="actions"><button id="importBtn">导入模板</button><button class="secondary" id="loadBtn">加载 Profile</button></div>
</section>

<section>
<h2>2. DeepSeek</h2>
<div class="row">
  <div><label>API Key</label><input id="apiKey" type="password" value="__DEFAULT_API_KEY__"></div>
  <div><label>模型</label><select id="model"><option>deepseek-v4-flash</option><option>deepseek-v4-pro</option></select></div>
  <div><label>生成方式</label><select id="generateMissing"><option value="false">只使用已填内容</option><option value="true">构建时生成未填章节</option></select></div>
</div>
</section>

<section id="formSection" hidden>
<h2>3. 局部内容</h2>
<div id="globalFields"></div>
<div id="fields"></div>
<div id="images"></div>
<div id="sections"></div>
<div class="actions">
  <button id="buildBtn">生成 Word</button>
  <button class="secondary" id="buildQaBtn">生成并审核</button>
  <a class="button ghost" id="download" href="#" hidden>下载 final.docx</a>
</div>
</section>

<section><h2>状态</h2><div class="status" id="status">等待导入。</div></section>
</main>
<script>
let activeProfile = null;

function status(text) { document.querySelector("#status").textContent = text; }
function el(tag, attrs = {}, children = []) {
  const node = document.createElement(tag);
  for (const [key, value] of Object.entries(attrs)) {
    if (key === "text") node.textContent = value;
    else if (key === "html") node.innerHTML = value;
    else node.setAttribute(key, value);
  }
  for (const child of children) node.appendChild(child);
  return node;
}
function fileToPayload(file) {
  return new Promise((resolve, reject) => {
    if (!file) return resolve(null);
    const reader = new FileReader();
    reader.onerror = reject;
    reader.onload = () => resolve({ name: file.name, data: String(reader.result).split(",")[1] || "" });
    reader.readAsDataURL(file);
  });
}
async function filesToPayload(input) {
  const files = Array.from(input?.files || []);
  const payloads = await Promise.all(files.map(fileToPayload));
  return payloads.filter(Boolean);
}
async function postJson(url, payload) {
  const res = await fetch(url, { method: "POST", headers: { "content-type": "application/json" }, body: JSON.stringify(payload) });
  const data = await res.json();
  if (!res.ok) throw new Error(data.error || res.statusText);
  return data;
}
async function getJson(url) {
  const res = await fetch(url);
  const data = await res.json();
  if (!res.ok) throw new Error(data.error || res.statusText);
  return data;
}
function fieldValue(path) {
  const input = document.querySelector(`[data-field-path="${CSS.escape(path)}"]`);
  return input ? input.value : "";
}
function collectProject() {
  const project = { project_id: document.querySelector("#projectId").value || "demo", profile_id: activeProfile.profile_id, global: {}, fields: {}, sections: {}, images: {} };
  document.querySelectorAll("[data-global-path]").forEach(input => {
    setDotted(project, input.dataset.globalPath, input.value);
  });
  document.querySelectorAll("[data-field-path]").forEach(input => {
    setDotted(project, input.dataset.fieldPath, input.value);
  });
  document.querySelectorAll("[data-section-id]").forEach(card => {
    const id = card.dataset.sectionId;
    setDotted(project, `sections.${id}.prompt`, card.querySelector("[data-section-prompt]").value);
    setDotted(project, `sections.${id}.content`, card.querySelector("[data-section-content]").value);
  });
  return project;
}
async function collectProjectWithFiles() {
  const project = collectProject();
  for (const card of document.querySelectorAll("[data-image-id]")) {
    const id = card.dataset.imageId;
    const mode = card.querySelector("[data-image-mode]").value;
    const file = card.querySelector("[data-image-file]").files[0];
    project.images[id] = {
      mode,
      file: mode === "replace" ? await fileToPayload(file) : null,
    };
  }
  return project;
}
function setDotted(obj, path, value) {
  const parts = path.split(".");
  let current = obj;
  for (const part of parts.slice(0, -1)) current = current[part] ||= {};
  current[parts[parts.length - 1]] = value;
}
function renderGlobalFields(profile) {
  const global = document.querySelector("#globalFields");
  global.innerHTML = "";
  global.appendChild(el("h3", { text: "全局信息" }));
  global.appendChild(el("p", { class: "muted", text: "这些内容会作为所有章节生成的共同背景，不直接改变模板版式。" }));
  const grid = el("div", { class: "global-grid" });
  for (const field of (profile.global_fields || [])) {
    const wrap = el("div");
    wrap.appendChild(el("label", { text: field.label || field.path || field.id }));
    const input = el("textarea", {
      "data-global-path": field.path || `global.${field.id}`,
      placeholder: "填写新项目的全局信息、资料或写作要求",
    });
    input.value = field.default || "";
    wrap.appendChild(input);
    grid.appendChild(wrap);
  }
  global.appendChild(grid);
  const actions = el("div", { class: "actions" });
  const generateAll = el("button", { type: "button", text: "根据全局信息整体生成" });
  generateAll.id = "generateAllBtn";
  generateAll.addEventListener("click", () => generateAllSections());
  actions.appendChild(generateAll);
  global.appendChild(actions);
  global.appendChild(el("div", { class: "global-status", id: "globalGenerateStatus", text: "" }));
}
function templateImagePreviewUrl(image) {
  return `/api/template-image?profileId=${encodeURIComponent(activeProfile.profile_id)}&imageId=${encodeURIComponent(image.id)}`;
}
function canPreviewImage(nameOrType = "") {
  const value = String(nameOrType || "").toLowerCase();
  return value.includes("image/png") || value.includes("image/jpeg") || value.includes("image/jpg") ||
    value.includes("image/gif") || value.includes("image/webp") || value.includes("image/svg") ||
    value.includes("image/x-wmf") || value.includes("image/wmf") || value.includes("image/emf") ||
    value.endsWith(".png") || value.endsWith(".jpg") || value.endsWith(".jpeg") ||
    value.endsWith(".gif") || value.endsWith(".webp") || value.endsWith(".svg") ||
    value.endsWith(".wmf") || value.endsWith(".emf");
}
function setPreviewFrame(frame, src, alt, fallbackText) {
  frame.innerHTML = "";
  if (src) {
    frame.appendChild(el("img", { src, alt }));
  } else {
    frame.appendChild(el("span", { text: fallbackText || "此图片格式浏览器无法直接预览" }));
  }
}
function showTemplatePreview(frame, image) {
  if (canPreviewImage(`${image.content_type || ""} ${image.partname || ""}`)) {
    setPreviewFrame(frame, templateImagePreviewUrl(image), image.label || image.id);
  } else {
    setPreviewFrame(frame, "", image.label || image.id, `${image.partname || "模板图片"}\n${image.content_type || "该格式不可直接预览"}`);
  }
}
function showUploadPreview(frame, image, file) {
  if (!file) {
    showTemplatePreview(frame, image);
    return;
  }
  if (canPreviewImage(`${file.type || ""} ${file.name || ""}`)) {
    setPreviewFrame(frame, URL.createObjectURL(file), file.name || "上传图片");
  } else {
    setPreviewFrame(frame, "", file.name || "上传图片", `已选择：${file.name || "上传图片"}\n此格式浏览器无法直接预览`);
  }
}
function renderImages(profile) {
  const root = document.querySelector("#images");
  root.innerHTML = "";
  const images = profile.targets?.images || [];
  if (!images.length) return;
  root.appendChild(el("h3", { text: "图片处理" }));
  root.appendChild(el("p", { class: "muted", text: "模板检测到图片。默认保留原图片；如需替换，请选择“上传新图片替换”。右侧可同时查看模板图和将要写入的新图。" }));
  for (const image of images) {
    const card = el("div", { class: "image-card", "data-image-id": image.id });
    const layout = el("div", { class: "image-layout" });
    const left = el("div");
    left.appendChild(el("strong", { text: image.label || image.id }));
    left.appendChild(el("div", { class: "muted", text: `${image.partname || ""}${image.content_type ? "；" + image.content_type : ""}` }));
    const select = el("select", { "data-image-mode": "" });
    select.appendChild(el("option", { value: "keep", text: "保留原图片" }));
    select.appendChild(el("option", { value: "replace", text: "上传新图片替换" }));
    const file = el("input", { type: "file", accept: "image/*", "data-image-file": "" });
    file.disabled = true;
    left.appendChild(el("label", { text: "处理方式" }));
    left.appendChild(select);
    left.appendChild(el("label", { text: "新图片" }));
    left.appendChild(file);

    const right = el("div", { class: "image-preview-grid" });
    const templateBox = el("div", { class: "preview-box" });
    templateBox.appendChild(el("div", { class: "preview-title", text: "模板图片预览" }));
    const templateFrame = el("div", { class: "preview-frame" });
    templateBox.appendChild(templateFrame);
    const uploadBox = el("div", { class: "preview-box" });
    uploadBox.appendChild(el("div", { class: "preview-title", text: "新图片预览" }));
    const uploadFrame = el("div", { class: "preview-frame" });
    uploadBox.appendChild(uploadFrame);
    right.appendChild(templateBox);
    right.appendChild(uploadBox);
    showTemplatePreview(templateFrame, image);
    showUploadPreview(uploadFrame, image, null);

    const updateUploadPreview = () => {
      showUploadPreview(uploadFrame, image, file.files[0] || null);
    };
    select.addEventListener("change", () => {
      file.disabled = select.value !== "replace";
      updateUploadPreview();
    });
    file.addEventListener("change", updateUploadPreview);

    layout.appendChild(left);
    layout.appendChild(right);
    card.appendChild(layout);
    root.appendChild(card);
  }
}
function renderProfile(profile) {
  activeProfile = profile;
  document.querySelector("#profileId").value = profile.profile_id;
  document.querySelector("#formSection").hidden = false;
  renderGlobalFields(profile);
  renderImages(profile);
  const fields = document.querySelector("#fields");
  fields.innerHTML = "";
  fields.appendChild(el("h3", { text: "项目字段 / 占位字段" }));
  const projectFields = profile.project_fields || {};
  for (const [path, spec] of Object.entries(projectFields)) {
    const wrap = el("div");
    wrap.appendChild(el("label", { text: `${spec.label || path} (${path})` }));
    wrap.appendChild(el("input", { "data-field-path": path, value: spec.default || "" }));
    fields.appendChild(wrap);
  }
  const sections = document.querySelector("#sections");
  sections.innerHTML = "";
  sections.appendChild(el("h3", { text: "章节/局部段落" }));
  for (const section of (profile.targets?.sections || [])) {
    const card = el("section", { "data-section-id": section.id });
    const left = el("div");
    left.appendChild(el("h3", { text: section.title || section.id }));
    left.appendChild(el("label", { text: "提示词" }));
    left.appendChild(el("textarea", { "data-section-prompt": "", placeholder: "写新项目资料、写作要求或要点" }));
    left.appendChild(el("label", { text: "生成/写回正文" }));
    left.appendChild(el("textarea", { "data-section-content": "", placeholder: "可手写，也可点生成覆盖这里" }));
    const fileWrap = el("div", { class: "file-row" });
    fileWrap.appendChild(el("label", { text: "上传参考资料（可选）" }));
    fileWrap.appendChild(el("input", { type: "file", multiple: "multiple", "data-section-files": "" }));
    left.appendChild(fileWrap);
    const actions = el("div", { class: "actions" });
    const gen = el("button", { type: "button", text: "生成本节" });
    gen.dataset.role = "generate";
    gen.addEventListener("click", () => generateSection(section, card, "generate"));
    const regen = el("button", { type: "button", class: "secondary", text: "重新生成" });
    regen.dataset.role = "regenerate";
    regen.addEventListener("click", () => generateSection(section, card, "regenerate"));
    actions.appendChild(gen);
    actions.appendChild(regen);
    left.appendChild(actions);
    left.appendChild(el("div", { class: "section-status", "data-section-status": "", text: "" }));
    const right = el("div");
    right.appendChild(el("strong", { text: "模板样例" }));
    right.appendChild(el("div", { class: "sample", text: section.template_sample || section.context_sample || "模板只检测到标题，未发现可摘取正文。可参考相邻章节结构填写提示词。" }));
    card.appendChild(el("div", { class: "target" }, [left, right]));
    sections.appendChild(card);
  }
  status(`已加载 profile：${profile.profile_id}\n目标：${JSON.stringify(profile.stats, null, 2)}`);
}
async function loadProfile() {
  const id = document.querySelector("#profileId").value.trim();
  if (!id) throw new Error("请先输入 Profile ID");
  renderProfile(await getJson(`/api/profile?id=${encodeURIComponent(id)}`));
}
async function generateSection(section, card) {
  const prompt = card.querySelector("[data-section-prompt]").value.trim();
  if (!prompt) { status("请先填写本节提示词。"); return; }
  status(`正在生成：${section.title}`);
  const data = await postJson("/api/generate-section", {
    profileId: activeProfile.profile_id,
    sectionId: section.id,
    prompt,
    project: collectProject(),
    apiKey: document.querySelector("#apiKey").value,
    model: document.querySelector("#model").value,
  });
  card.querySelector("[data-section-content]").value = data.text || "";
  status(`已生成：${section.title}`);
}
function setSectionBusy(card, busy) {
  card.querySelectorAll("button").forEach(button => {
    if (button.dataset.role === "generate" || button.dataset.role === "regenerate") {
      button.disabled = busy;
      button.textContent = busy && button.dataset.role === "generate" ? "生成中..." :
        button.dataset.role === "generate" ? "生成本节" :
        busy && button.dataset.role === "regenerate" ? "生成中..." : "重新生成";
    }
  });
}
function setSectionStatus(card, text, isError = false) {
  const target = card.querySelector("[data-section-status]");
  if (!target) return;
  target.textContent = text || "";
  target.classList.toggle("error", Boolean(isError));
}
async function generateSection(section, card, mode = "generate") {
  const prompt = card.querySelector("[data-section-prompt]").value.trim();
  setSectionBusy(card, true);
  setSectionStatus(card, mode === "regenerate" ? "重新生成中..." : "生成中...");
  status(`${mode === "regenerate" ? "重新生成中" : "生成中"}：${section.title || section.id}`);
  try {
    const data = await postJson("/api/generate-section", {
      profileId: activeProfile.profile_id,
      sectionId: section.id,
      prompt,
      mode,
      previousText: card.querySelector("[data-section-content]").value || "",
      attachments: await filesToPayload(card.querySelector("[data-section-files]")),
      project: collectProject(),
      apiKey: document.querySelector("#apiKey").value,
      model: document.querySelector("#model").value,
    });
    card.querySelector("[data-section-content]").value = data.text || "";
    setSectionStatus(card, mode === "regenerate" ? "已重新生成" : "已生成");
    status(`已生成：${section.title || section.id}`);
    return true;
  } catch (err) {
    setSectionStatus(card, `生成失败：${err.message}`, true);
    status(`生成失败：${err.message}`);
    return false;
  } finally {
    setSectionBusy(card, false);
  }
}
function setGlobalStatus(text, isError = false) {
  const target = document.querySelector("#globalGenerateStatus");
  if (!target) return;
  target.textContent = text || "";
  target.classList.toggle("error", Boolean(isError));
}
function setGlobalBusy(busy) {
  const button = document.querySelector("#generateAllBtn");
  if (!button) return;
  button.disabled = busy;
  button.textContent = busy ? "整体生成中..." : "根据全局信息整体生成";
}
async function generateAllSections() {
  const cards = Array.from(document.querySelectorAll("[data-section-id]"));
  if (!activeProfile || cards.length === 0) {
    setGlobalStatus("请先导入或加载模板。", true);
    status("请先导入或加载模板。");
    return;
  }
  setGlobalBusy(true);
  setGlobalStatus(`整体生成中：0 / ${cards.length}`);
  status(`整体生成中：0 / ${cards.length}`);
  let completed = 0;
  try {
    for (const card of cards) {
      const section = (activeProfile.targets?.sections || []).find(item => String(item.id) === String(card.dataset.sectionId));
      if (!section) continue;
      setGlobalStatus(`整体生成中：${completed + 1} / ${cards.length}\n当前：${section.title || section.id}`);
      const ok = await generateSection(section, card, "generate");
      if (!ok) {
        throw new Error(`${section.title || section.id} 生成失败`);
      }
      completed += 1;
    }
    setGlobalStatus(`整体生成完成：${completed} / ${cards.length}`);
    status(`整体生成完成：${completed} / ${cards.length}`);
  } catch (err) {
    setGlobalStatus(`整体生成失败：${err.message}`, true);
    status(`整体生成失败：${err.message}`);
  } finally {
    setGlobalBusy(false);
  }
}
document.querySelector("#importBtn").addEventListener("click", async () => {
  try {
    status("导入中...");
    const file = document.querySelector("#templateFile").files[0];
    const payload = {
      templatePath: document.querySelector("#templatePath").value,
      templateFile: await fileToPayload(file),
      profileId: document.querySelector("#profileId").value,
      projectId: document.querySelector("#projectId").value,
    };
    renderProfile(await postJson("/api/import", payload));
  } catch (err) { status(`导入失败：${err.message}`); }
});
document.querySelector("#loadBtn").addEventListener("click", async () => {
  try { await loadProfile(); } catch (err) { status(`加载失败：${err.message}`); }
});
async function build(qa) {
  status(qa ? "生成并审核中..." : "生成中...");
  const data = await postJson("/api/build", {
    profileId: activeProfile.profile_id,
    projectId: document.querySelector("#projectId").value || "demo",
    project: await collectProjectWithFiles(),
    generate: document.querySelector("#generateMissing").value === "true",
    apiKey: document.querySelector("#apiKey").value,
    model: document.querySelector("#model").value,
    qa,
  });
  const link = document.querySelector("#download");
  link.href = data.downloadUrl;
  link.setAttribute("download", "final.docx");
  link.hidden = false;
  status(`已生成 Word。\n本机文件：${data.final_docx}\n浏览器下载地址：${data.downloadUrl}`);
  setTimeout(() => link.click(), 50);
}
document.querySelector("#buildBtn").addEventListener("click", () => build(false).catch(err => status(`生成失败：${err.message}`)));
document.querySelector("#buildQaBtn").addEventListener("click", () => build(true).catch(err => status(`审核失败：${err.message}`)));
</script>
</body>
</html>
"""
    return template.replace("__DEFAULT_API_KEY__", escaped_key).encode("utf-8")


def _save_uploaded_template(root: Path, payload: dict[str, Any]) -> Path | None:
    file_payload = payload.get("templateFile")
    if not isinstance(file_payload, dict):
        return None
    name = Path(str(file_payload.get("name") or "template.docx")).name
    encoded = str(file_payload.get("data") or "")
    if not encoded:
        return None
    uploads = root / "uploads"
    uploads.mkdir(parents=True, exist_ok=True)
    target = uploads / name
    target.write_bytes(base64.b64decode(encoded))
    return target


def _clip_text(text: str, limit: int = 12000) -> str:
    text = text.strip()
    if len(text) <= limit:
        return text
    return text[:limit] + "\n...[truncated]"


def _decode_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_attachment_texts(uploads: object) -> list[dict[str, Any]]:
    if not isinstance(uploads, list):
        return []

    extracted: list[dict[str, Any]] = []
    text_suffixes = {".txt", ".md", ".csv", ".json", ".yaml", ".yml", ".log"}
    for item in uploads:
        if not isinstance(item, dict):
            continue
        name = Path(str(item.get("name") or "attachment")).name
        encoded = str(item.get("data") or "")
        if not encoded:
            continue
        try:
            data = base64.b64decode(encoded)
        except Exception as exc:
            extracted.append({"name": name, "error": f"base64 decode failed: {exc}"})
            continue

        suffix = Path(name).suffix.lower()
        text = ""
        error = ""
        try:
            if suffix in text_suffixes:
                text = _decode_text_bytes(data)
            elif suffix == ".docx":
                from docx import Document

                document = Document(BytesIO(data))
                chunks: list[str] = []
                chunks.extend(p.text for p in document.paragraphs if p.text.strip())
                for table in document.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        if any(cells):
                            chunks.append(" | ".join(cells))
                text = "\n".join(chunks)
            elif suffix == ".pdf":
                from pypdf import PdfReader

                reader = PdfReader(BytesIO(data))
                pages = []
                for page in reader.pages[:20]:
                    pages.append(page.extract_text() or "")
                text = "\n".join(pages)
            else:
                error = "unsupported attachment type for text extraction"
        except Exception as exc:
            error = f"text extraction failed: {exc}"

        record: dict[str, Any] = {
            "name": name,
            "size_bytes": len(data),
        }
        if text.strip():
            record["text_excerpt"] = _clip_text(text)
        if error:
            record["note"] = error
        extracted.append(record)
    return extracted


def _project_download(root: Path, project_id: str) -> Path:
    return (root / "projects" / project_id / "final.docx").resolve()


def _safe_preview_name(value: str) -> str:
    cleaned = "".join(ch if ch.isalnum() or ch in "-_" else "_" for ch in value)
    return cleaned[:80] or "image"


def _previewable_image_bytes(
    profile_dir: Path,
    image_id: str,
    entry: str,
    data: bytes,
    content_type: str,
) -> tuple[bytes, str]:
    suffix = Path(entry).suffix.lower()
    if suffix not in {".wmf", ".emf"} and content_type not in {"image/x-wmf", "image/wmf", "image/emf"}:
        return data, content_type

    soffice = shutil.which("soffice.com") or shutil.which("soffice")
    if not soffice:
        return data, content_type

    preview_dir = profile_dir / "_preview_cache"
    preview_dir.mkdir(parents=True, exist_ok=True)
    safe_name = _safe_preview_name(image_id)
    source = preview_dir / f"{safe_name}{suffix or '.wmf'}"
    target = preview_dir / f"{safe_name}.png"
    if target.exists() and target.stat().st_size > 0:
        return target.read_bytes(), "image/png"

    source.write_bytes(data)
    with tempfile.TemporaryDirectory(prefix="template_to_doc_v2_lo_preview_") as lo_profile:
        completed = subprocess.run(
            [
                soffice,
                f"-env:UserInstallation={Path(lo_profile).resolve().as_uri()}",
                "--headless",
                "--convert-to",
                "png",
                "--outdir",
                str(preview_dir),
                str(source),
            ],
            text=True,
            capture_output=True,
            timeout=60,
            check=False,
        )
    if completed.returncode != 0 or not target.exists():
        return data, content_type
    return target.read_bytes(), "image/png"


def _template_image_bytes(
    root: Path,
    pipeline: TemplateToDocV2Pipeline,
    profile_id: str,
    image_id: str,
) -> tuple[bytes, str]:
    profile = pipeline.load_profile(profile_id)
    image = next(
        (
            item
            for item in profile.get("targets", {}).get("images", [])
            if str(item.get("id")) == image_id
        ),
        None,
    )
    if not image:
        raise FileNotFoundError(f"image not found: {image_id}")
    entry = str(image.get("partname") or "").lstrip("/")
    if not entry.startswith("word/media/"):
        raise ValueError("template image path is not a Word media entry")
    reference = pipeline.profile_reference(profile_id).resolve()
    if root not in reference.parents:
        raise ValueError("template reference is outside the project root")
    with zipfile.ZipFile(reference) as archive:
        data = archive.read(entry)
    content_type = str(image.get("content_type") or "application/octet-stream")
    return _previewable_image_bytes(
        pipeline.paths.profile_dir(profile_id),
        image_id,
        entry,
        data,
        content_type,
    )


def serve(root: str | Path, host: str = "127.0.0.1", port: int = 8770) -> None:
    base = Path(root).resolve()
    pipeline = TemplateToDocV2Pipeline(base)

    class Handler(BaseHTTPRequestHandler):
        def _send(self, status: int, body: bytes, content_type: str) -> None:
            self.send_response(status)
            self.send_header("content-type", content_type)
            self.send_header("content-length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)

        def _json(self, status: int, data: dict[str, Any]) -> None:
            self._send(status, _json_bytes(data), "application/json; charset=utf-8")

        def do_GET(self) -> None:  # noqa: N802
            parsed = urlparse(self.path)
            if parsed.path == "/":
                self._send(
                    200,
                    _html(_read_default_deepseek_key(base)),
                    "text/html; charset=utf-8",
                )
                return
            if parsed.path == "/api/profile":
                try:
                    profile_id = parse_qs(parsed.query).get("id", [""])[0]
                    self._json(200, pipeline.load_profile(profile_id))
                except Exception as exc:
                    self._json(404, {"error": str(exc)})
                return
            if parsed.path == "/api/template-image":
                try:
                    query = parse_qs(parsed.query)
                    data, content_type = _template_image_bytes(
                        base,
                        pipeline,
                        query.get("profileId", [""])[0],
                        query.get("imageId", [""])[0],
                    )
                    self._send(200, data, content_type)
                except Exception as exc:
                    self._json(404, {"error": str(exc)})
                return
            if parsed.path.startswith("/download/") and parsed.path.endswith("/final.docx"):
                project_id = unquote(parsed.path.split("/")[2])
                path = _project_download(base, project_id)
                if not path.exists() or base not in path.parents:
                    self._json(404, {"error": "final.docx not found"})
                    return
                self._send(200, path.read_bytes(), DOCX_CONTENT_TYPE)
                return
            self._json(404, {"error": "not found"})

        def do_POST(self) -> None:  # noqa: N802
            try:
                payload = _read_json(self)
                if self.path == "/api/import":
                    uploaded = _save_uploaded_template(base, payload)
                    template = uploaded or Path(str(payload.get("templatePath") or ""))
                    result = pipeline.import_template(
                        template,
                        profile_id=str(payload.get("profileId") or "") or None,
                        project_id=str(payload.get("projectId") or "") or None,
                    )
                    self._json(200, result)
                    return
                if self.path == "/api/generate-section":
                    profile_id = str(payload.get("profileId") or "")
                    section_id = str(payload.get("sectionId") or "")
                    profile = pipeline.load_profile(profile_id)
                    section = next(
                        (
                            item
                            for item in profile.get("targets", {}).get("sections", [])
                            if str(item.get("id")) == section_id
                        ),
                        None,
                    )
                    if not section:
                        raise ValueError(f"section not found: {section_id}")
                    provider = deepseek_provider(
                        str(payload.get("apiKey") or ""),
                        str(payload.get("model") or "deepseek-v4-flash"),
                    )
                    project_context = dict(payload.get("project") or {})
                    project_context["_generation"] = {
                        "mode": str(payload.get("mode") or "generate"),
                        "previous_text": str(payload.get("previousText") or ""),
                    }
                    text = generate_local_text(
                        prompt=str(payload.get("prompt") or ""),
                        template_sample=str(section.get("template_sample") or ""),
                        target_label=str(section.get("title") or section_id),
                        project_context=project_context,
                        attachments=_extract_attachment_texts(payload.get("attachments")),
                        provider=provider,
                    )
                    self._json(200, {"text": text})
                    return
                if self.path == "/api/build":
                    profile_id = str(payload.get("profileId") or "")
                    project_id = str(payload.get("projectId") or "demo")
                    project = dict(payload.get("project") or {})
                    output = pipeline.build(
                        profile_id,
                        project_id,
                        project=project,
                        generate=bool(payload.get("generate")),
                        api_key=str(payload.get("apiKey") or ""),
                        model=str(payload.get("model") or "deepseek-v4-flash"),
                        qa=bool(payload.get("qa")),
                    )
                    output["downloadUrl"] = f"/download/{quote(project_id)}/final.docx"
                    self._json(200, output)
                    return
            except Exception as exc:
                self._json(500, {"error": str(exc)})
                return
            self._json(404, {"error": "not found"})

        def log_message(self, format: str, *args: Any) -> None:  # noqa: A003
            return

    ThreadingHTTPServer((host, port), Handler).serve_forever()
